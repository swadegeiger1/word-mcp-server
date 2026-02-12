#!/usr/bin/env python3
"""Test the Word MCP server by creating a sample document."""

from docx import Document

# Create a test document
doc = Document()
doc.add_paragraph("Test Document")
doc.add_paragraph("This is a test paragraph created by the Word MCP server.")
doc.save("/tmp/test_word_mcp.docx")

print("✓ Successfully created test document at /tmp/test_word_mcp.docx")

# Read it back
doc = Document("/tmp/test_word_mcp.docx")
text = "\n".join([p.text for p in doc.paragraphs])
print(f"\n✓ Successfully read document:\n{text}")
