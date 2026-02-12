#!/usr/bin/env python3
"""Minimal MCP server for Word document operations using python-docx."""

from mcp.server import Server
from mcp.types import Tool, TextContent
from docx import Document
from docx.shared import Pt, RGBColor
import mcp.server.stdio
import json

app = Server("word-mcp")

@app.list_tools()
async def list_tools() -> list[Tool]:
    return [
        Tool(
            name="read_word_document",
            description="Read text content from a Word document",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Path to the .docx file"}
                },
                "required": ["file_path"]
            }
        ),
        Tool(
            name="create_word_document",
            description="Create a new Word document with text content",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Path where to save the .docx file"},
                    "content": {"type": "string", "description": "Text content to write"}
                },
                "required": ["file_path", "content"]
            }
        ),
        Tool(
            name="append_to_document",
            description="Append text to an existing Word document",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Path to the .docx file"},
                    "content": {"type": "string", "description": "Text to append"}
                },
                "required": ["file_path", "content"]
            }
        )
    ]

@app.call_tool()
async def call_tool(name: str, arguments: dict) -> list[TextContent]:
    if name == "read_word_document":
        doc = Document(arguments["file_path"])
        text = "\n".join([p.text for p in doc.paragraphs])
        return [TextContent(type="text", text=text)]
    
    elif name == "create_word_document":
        doc = Document()
        for line in arguments["content"].split("\n"):
            doc.add_paragraph(line)
        doc.save(arguments["file_path"])
        return [TextContent(type="text", text=f"Created document: {arguments['file_path']}")]
    
    elif name == "append_to_document":
        doc = Document(arguments["file_path"])
        for line in arguments["content"].split("\n"):
            doc.add_paragraph(line)
        doc.save(arguments["file_path"])
        return [TextContent(type="text", text=f"Appended to: {arguments['file_path']}")]
    
    raise ValueError(f"Unknown tool: {name}")

async def main():
    async with mcp.server.stdio.stdio_server() as (read_stream, write_stream):
        await app.run(read_stream, write_stream, app.create_initialization_options())

if __name__ == "__main__":
    import asyncio
    asyncio.run(main())
